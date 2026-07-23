#!/usr/bin/env python3
"""
Vult een Word-sjabloon met content van een AI-agent.

Het sjabloon levert de stijl: logo, header/footer, marges, kleuren, stijlbladen.
De content (dict) levert alleen structuur en tekst. Dit script koppelt die twee.

Elk sjabloon heeft een eigen stijlconfiguratie (welke styleId hoort bij welk
bloktype). Bestaat een stijl niet in het document, dan valt het script terug
op een alternatief, zodat er nooit een crash ontstaat op een ontbrekende stijl.

Sjabloon bevat placeholders:
    {{titel}} {{jobnummer}} {{klant}} {{contactpersoon}} {{opdracht}} {{datum}}
    {{content}}   <- hier worden de contentblokken ingevoegd
"""

import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Twips

CONTENT_PLACEHOLDER = "{{content}}"

# Terugvalvolgorde per bloktype: bestaat de eerste stijl niet, dan de volgende.
# Laatste redmiddel is altijd Standaard.
TERUGVAL = {
    "titel":     ["Titel", "Kop1", "Standaard"],
    "kop1":      ["Kop1", "Kop2", "Standaard"],
    "kop2":      ["Kop2", "Kop1", "Standaard"],
    "kop3":      ["Kop3", "Kop2", "Standaard"],
    "accent":    ["Kop4", "Kop3", "Standaard"],
    "standaard": ["Standaard"],
    "citaat":    ["Citaat", "Plattetekst", "Standaard"],
    "bullet":    ["Lijstalinea", "Standaard"],
}


def beschikbare_stijlen(doc):
    """Alle styleId's die daadwerkelijk in het document voorkomen."""
    ids = set()
    for st in doc.styles.element:
        sid = st.get(qn("w:styleId"))
        if sid:
            ids.add(sid)
    return ids


def kies_stijl(soort, config, aanwezig):
    """Bepaal de styleId voor een bloktype, met terugval als hij ontbreekt."""
    kandidaten = []
    gewenst = (config.get("stijlen") or {}).get(soort)
    if gewenst:
        kandidaten.append(gewenst)
    kandidaten.extend(TERUGVAL.get(soort, []))
    for sid in kandidaten:
        if sid in aanwezig:
            return sid
    return "Standaard"


def vul_velden(doc, velden):
    """Vervangt {{sleutel}} door de waarde, in body, headers en footers."""
    bronnen = list(doc.paragraphs)
    for sectie in doc.sections:
        for deel in (sectie.header, sectie.footer,
                     sectie.first_page_header, sectie.first_page_footer):
            if deel is not None:
                bronnen.extend(deel.paragraphs)

    for alinea in bronnen:
        for run in alinea.runs:
            for sleutel, waarde in velden.items():
                merk = "{{" + sleutel + "}}"
                if merk in run.text:
                    run.text = run.text.replace(merk, str(waarde))


def ruim_lege_placeholders(doc):
    """Haalt placeholders weg die niet gevuld konden worden."""
    for alinea in list(doc.paragraphs):
        if not re.search(r"\{\{[a-z_]+\}\}", alinea.text):
            continue
        for run in alinea.runs:
            run.text = re.sub(r"\{\{[a-z_]+\}\}", "", run.text)


def zoek_content_alinea(doc):
    for alinea in doc.paragraphs:
        if CONTENT_PLACEHOLDER in alinea.text:
            return alinea
    raise ValueError(f"{CONTENT_PLACEHOLDER} niet gevonden in het sjabloon")


def schrijf_tekst(alinea, tekst):
    """Plaatst tekst in een alinea en maakt **stukken** echt vet."""
    for deel in re.split(r"(\*\*[^*]+\*\*)", tekst):
        if not deel:
            continue
        if deel.startswith("**") and deel.endswith("**"):
            run = alinea.add_run(deel[2:-2])
            run.bold = True
        else:
            alinea.add_run(deel)


def zet_stijl(alinea, style_id):
    pPr = alinea._p.get_or_add_pPr()
    for oud in pPr.findall(qn("w:pStyle")):
        pPr.remove(oud)
    el = OxmlElement("w:pStyle")
    el.set(qn("w:val"), style_id)
    pPr.insert(0, el)


def maak_bullet(alinea, num_id):
    pPr = alinea._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def zet_tabelstijl(tabel, style_id):
    tblPr = tabel._tbl.tblPr
    for oud in tblPr.findall(qn("w:tblStyle")):
        tblPr.remove(oud)
    el = OxmlElement("w:tblStyle")
    el.set(qn("w:val"), style_id)
    tblPr.insert(0, el)

    for oud in tblPr.findall(qn("w:tblLook")):
        tblPr.remove(oud)
    look = OxmlElement("w:tblLook")
    for k, v in [("firstRow", "1"), ("lastRow", "0"), ("firstColumn", "0"),
                 ("lastColumn", "0"), ("noHBand", "0"), ("noVBand", "1")]:
        look.set(qn("w:" + k), v)
    look.set(qn("w:val"), "04A0")
    tblPr.append(look)


def bouw_tabel(doc, kop, rijen, breedtes, config, aanwezig):
    kolommen = len(kop)
    breedte_totaal = config.get("paginabreedte", 9072)
    if not breedtes:
        breedtes = [breedte_totaal // kolommen] * kolommen
        breedtes[-1] += breedte_totaal - sum(breedtes)

    tabel = doc.add_table(rows=1 + len(rijen), cols=kolommen)

    for sid in [config.get("tabelstijl"), "Tabelraster", "Standaardtabel"]:
        if sid and sid in aanwezig:
            zet_tabelstijl(tabel, sid)
            break

    koprij = config.get("koprijstijl")
    koprij = koprij if koprij in aanwezig else None

    for i, tekst in enumerate(kop):
        cel = tabel.cell(0, i)
        cel.text = str(tekst)
        cel.width = Twips(breedtes[i])
        if koprij:
            zet_stijl(cel.paragraphs[0], koprij)
    for r, rij in enumerate(rijen, start=1):
        for i, tekst in enumerate(rij[:kolommen]):
            cel = tabel.cell(r, i)
            cel.text = "" if tekst is None else str(tekst)
            cel.width = Twips(breedtes[i])

    trPr = tabel.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))   # koprij herhalen bij paginaovergang
    return tabel


def bouw(sjabloon_pad, content, uitvoer, config=None):
    """
    sjabloon_pad : pad naar het .docx-sjabloon
    content      : dict met 'velden' en 'blokken'
    uitvoer      : bestandspad (str) OF een open bytes-buffer
    config       : dict met stijlen/tabelstijl/bullet_numid voor dit sjabloon
    """
    config = config or {}
    doc = Document(sjabloon_pad)
    aanwezig = beschikbare_stijlen(doc)
    bullet_numid = config.get("bullet_numid", 1)

    vul_velden(doc, content.get("velden", {}))
    anker = zoek_content_alinea(doc)

    for blok in content.get("blokken", []):
        soort = (blok.get("type") or "standaard").lower()

        if soort == "tabel":
            tabel = bouw_tabel(doc, blok["kop"], blok.get("rijen", []),
                               blok.get("breedtes"), config, aanwezig)
            anker._p.addprevious(tabel._tbl)
            witregel = doc.add_paragraph()
            zet_stijl(witregel, kies_stijl("standaard", config, aanwezig))
            anker._p.addprevious(witregel._p)
            continue

        if soort not in TERUGVAL:
            soort = "standaard"

        alinea = doc.add_paragraph()
        schrijf_tekst(alinea, blok.get("tekst", ""))
        zet_stijl(alinea, kies_stijl(soort, config, aanwezig))
        if soort == "bullet":
            maak_bullet(alinea, bullet_numid)
        anker._p.addprevious(alinea._p)

    anker._p.getparent().remove(anker._p)
    ruim_lege_placeholders(doc)
    doc.save(uitvoer)
    return uitvoer
